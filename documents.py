"""Invoice arithmetic and document rendering (PDF + DOCX).

All currency math lives here, in Decimal with ROUND_HALF_UP. The LLM never
computes a number that reaches a client document — 0.1 + 0.2 != 0.3 in float,
and a client with a spreadsheet will find the cent.
"""

import io
import re
from datetime import date
from decimal import Decimal, ROUND_HALF_UP

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)

from config import CURRENCY_SYMBOL
from storage import pretty_date

TWO_PLACES = Decimal("0.01")

INK = colors.HexColor("#111418")
MUTED = colors.HexColor("#6b7280")
RULE = colors.HexColor("#d9dde3")
ACCENT = colors.HexColor("#1f3a5f")
BAND = colors.HexColor("#f3f5f7")


# ─────────────────────────── money ───────────────────────────

def D(value) -> Decimal:
    """Coerce anything to Decimal without going through float."""
    if isinstance(value, Decimal):
        return value
    if value is None or value == "":
        return Decimal("0")
    return Decimal(str(value))


def money(value) -> Decimal:
    return D(value).quantize(TWO_PLACES, rounding=ROUND_HALF_UP)


def fmt(value) -> str:
    """1800 -> '$1,800.00'. Always symbol, always two decimals."""
    return f"{CURRENCY_SYMBOL}{money(value):,.2f}"


def fmt_hours(value) -> str:
    """18.0 -> '18', 17.5 -> '17.5'. Trailing .0 on hours is noise."""
    d = D(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return str(d.normalize()) if d == d.to_integral_value() else f"{d:g}"


def compute_invoice(line_items, tax_rate=0):
    """The single source of arithmetic truth.

    Returns (items, subtotal, tax_amount, total) with every value quantized
    to 2 places at each step — not once at the end, which would let rounding
    error accumulate across many line items.
    """
    computed = []
    subtotal = Decimal("0.00")

    for raw in line_items or []:
        hours = D(raw.get("hours", 0))
        rate = money(raw.get("rate", 0))
        line_total = money(hours * rate)
        subtotal += line_total
        computed.append(
            {
                "description": (raw.get("description") or "Work").strip(),
                "hours": hours,
                "rate": rate,
                "subtotal": line_total,
            }
        )

    subtotal = money(subtotal)
    tax_amount = money(subtotal * (D(tax_rate) / Decimal("100")))
    total = money(subtotal + tax_amount)
    return computed, subtotal, tax_amount, total


# ─────────────────────────── markdown -> flowables ───────────────────────────

def _clean_inline(text: str) -> str:
    """Convert **bold**/*italic* to ReportLab tags and escape the rest."""
    text = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", text)
    return text


# ─────────────────────────── proposal PDF ───────────────────────────

def proposal_pdf(body_md: str, client_name: str, project_title: str, freelancer: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=f"Proposal — {project_title}", author=freelancer.get("name", ""),
    )

    base = getSampleStyleSheet()
    st_name = ParagraphStyle("Name", parent=base["Normal"], fontName="Helvetica-Bold",
                             fontSize=15, textColor=INK, spaceAfter=2, leading=18)
    st_contact = ParagraphStyle("Contact", parent=base["Normal"], fontName="Helvetica",
                                fontSize=8.5, textColor=MUTED, spaceAfter=0, leading=12)
    st_title = ParagraphStyle("Title", parent=base["Normal"], fontName="Helvetica-Bold",
                              fontSize=19, textColor=INK, spaceBefore=14, spaceAfter=3, leading=23)
    st_sub = ParagraphStyle("Sub", parent=base["Normal"], fontName="Helvetica",
                            fontSize=10, textColor=MUTED, spaceAfter=16, leading=14)
    st_h2 = ParagraphStyle("H2", parent=base["Normal"], fontName="Helvetica-Bold",
                           fontSize=11.5, textColor=ACCENT, spaceBefore=15, spaceAfter=5, leading=15)
    st_body = ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica",
                             fontSize=10, textColor=INK, leading=15.5, spaceAfter=7)
    st_bullet = ParagraphStyle("Bullet", parent=st_body, leftIndent=13,
                               bulletIndent=3, spaceAfter=3.5)

    flow = []
    name = freelancer.get("name") or "Freelance Professional"
    contact = " · ".join(x for x in [freelancer.get("email"), freelancer.get("phone")] if x)

    flow.append(Paragraph(_clean_inline(name), st_name))
    if contact:
        flow.append(Paragraph(_clean_inline(contact), st_contact))
    flow.append(Spacer(1, 7))
    flow.append(Table([[""]], colWidths=[doc.width],
                      style=TableStyle([("LINEBELOW", (0, 0), (-1, -1), 1, ACCENT)])))

    flow.append(Paragraph(_clean_inline(project_title or "Project Proposal"), st_title))
    flow.append(Paragraph(
        f"Prepared for {_clean_inline(client_name)}  ·  {pretty_date(date.today())}", st_sub))

    for raw_line in (body_md or "").split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()

        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            if heading:
                flow.append(Paragraph(_clean_inline(heading), st_h2))
        elif stripped.startswith(("- ", "* ", "• ")):
            flow.append(Paragraph(_clean_inline(stripped[2:].strip()), st_bullet,
                                  bulletText="•"))
        elif re.match(r"^\d+[.)]\s+", stripped):
            marker, _, rest = stripped.partition(" ")
            flow.append(Paragraph(_clean_inline(rest.strip()), st_bullet, bulletText=marker))
        elif set(stripped) <= set("-–—=_") and len(stripped) >= 3:
            flow.append(Spacer(1, 4))
        else:
            flow.append(Paragraph(_clean_inline(stripped), st_body))

    doc.build(flow)
    return buf.getvalue()


# ─────────────────────────── proposal DOCX ───────────────────────────

def proposal_docx(body_md: str, client_name: str, project_title: str, freelancer: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    head = document.add_paragraph()
    run = head.add_run(freelancer.get("name") or "Freelance Professional")
    run.bold = True
    run.font.size = Pt(15)

    contact = " · ".join(x for x in [freelancer.get("email"), freelancer.get("phone")] if x)
    if contact:
        crun = document.add_paragraph().add_run(contact)
        crun.font.size = Pt(9)
        crun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    title = document.add_paragraph()
    trun = title.add_run(project_title or "Project Proposal")
    trun.bold = True
    trun.font.size = Pt(18)

    sub = document.add_paragraph()
    srun = sub.add_run(f"Prepared for {client_name}  ·  {pretty_date(date.today())}")
    srun.font.size = Pt(9.5)
    srun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def strip_md(text: str) -> str:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text or "")
        return re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", text)

    for raw_line in (body_md or "").split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            if heading:
                hp = document.add_paragraph()
                hrun = hp.add_run(strip_md(heading))
                hrun.bold = True
                hrun.font.size = Pt(12)
                hrun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        elif stripped.startswith(("- ", "* ", "• ")):
            document.add_paragraph(strip_md(stripped[2:].strip()), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(strip_md(stripped.split(" ", 1)[1].strip()),
                                   style="List Number")
        elif set(stripped) <= set("-–—=_") and len(stripped) >= 3:
            continue
        else:
            document.add_paragraph(strip_md(stripped))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ─────────────────────────── invoice PDF ───────────────────────────

def invoice_pdf(invoice: dict, client: dict, freelancer: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=f"Invoice #{invoice.get('number')}", author=freelancer.get("name", ""),
    )

    base = getSampleStyleSheet()
    st_word = ParagraphStyle("Word", parent=base["Normal"], fontName="Helvetica-Bold",
                             fontSize=24, textColor=ACCENT, leading=27)
    st_num = ParagraphStyle("Num", parent=base["Normal"], fontName="Helvetica",
                            fontSize=10, textColor=MUTED, alignment=TA_RIGHT, leading=15)
    st_label = ParagraphStyle("Label", parent=base["Normal"], fontName="Helvetica-Bold",
                              fontSize=7.5, textColor=MUTED, leading=11, spaceAfter=3)
    st_block = ParagraphStyle("Block", parent=base["Normal"], fontName="Helvetica",
                              fontSize=9.5, textColor=INK, leading=13.5)
    st_th = ParagraphStyle("TH", parent=base["Normal"], fontName="Helvetica-Bold",
                           fontSize=8.5, textColor=colors.white, leading=12)
    st_thr = ParagraphStyle("THR", parent=st_th, alignment=TA_RIGHT)
    st_td = ParagraphStyle("TD", parent=base["Normal"], fontName="Helvetica",
                           fontSize=9.5, textColor=INK, leading=13)
    st_tdr = ParagraphStyle("TDR", parent=st_td, alignment=TA_RIGHT)
    st_note = ParagraphStyle("Note", parent=base["Normal"], fontName="Helvetica",
                             fontSize=9, textColor=INK, leading=13)
    st_foot = ParagraphStyle("Foot", parent=base["Normal"], fontName="Helvetica",
                             fontSize=8, textColor=MUTED, leading=11, alignment=TA_CENTER)

    flow = []
    W = doc.width

    # Masthead
    meta_lines = [
        f"<b>Invoice #{invoice.get('number')}</b>",
        f"Issued {pretty_date(invoice.get('issue_date'))}",
        f"Due {pretty_date(invoice.get('due_date'))}",
    ]
    flow.append(Table(
        [[Paragraph("INVOICE", st_word), Paragraph("<br/>".join(meta_lines), st_num)]],
        colWidths=[W * 0.5, W * 0.5],
        style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ])))
    flow.append(Table([[""]], colWidths=[W],
                      style=TableStyle([("LINEBELOW", (0, 0), (-1, -1), 1.2, ACCENT)])))
    flow.append(Spacer(1, 12))

    # From / Bill to
    from_bits = [f"<b>{_clean_inline(freelancer.get('name') or '—')}</b>"]
    if freelancer.get("email"):
        from_bits.append(_clean_inline(freelancer["email"]))
    if freelancer.get("phone"):
        from_bits.append(_clean_inline(freelancer["phone"]))

    to_bits = [f"<b>{_clean_inline(client.get('name') or '—')}</b>"]
    if client.get("company") and client["company"] != client.get("name"):
        to_bits.append(_clean_inline(client["company"]))
    if client.get("email"):
        to_bits.append(_clean_inline(client["email"]))

    flow.append(Table(
        [[Paragraph("FROM", st_label), Paragraph("BILL TO", st_label)],
         [Paragraph("<br/>".join(from_bits), st_block),
          Paragraph("<br/>".join(to_bits), st_block)]],
        colWidths=[W * 0.5, W * 0.5],
        style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ])))

    if invoice.get("project"):
        flow.append(Spacer(1, 10))
        flow.append(Paragraph("PROJECT", st_label))
        flow.append(Paragraph(_clean_inline(invoice["project"]), st_block))

    flow.append(Spacer(1, 16))

    # Line items
    rows = [[Paragraph("DESCRIPTION", st_th), Paragraph("HOURS", st_thr),
             Paragraph("RATE", st_thr), Paragraph("AMOUNT", st_thr)]]
    for item in invoice.get("line_items", []):
        rows.append([
            Paragraph(_clean_inline(item.get("description", "")), st_td),
            Paragraph(fmt_hours(item.get("hours", 0)), st_tdr),
            Paragraph(fmt(item.get("rate", 0)), st_tdr),
            Paragraph(fmt(item.get("subtotal", 0)), st_tdr),
        ])

    col_widths = [W * 0.52, W * 0.13, W * 0.16, W * 0.19]
    table = Table(rows, colWidths=col_widths, repeatRows=1)
    table_style = [
        ("BACKGROUND", (0, 0), (-1, 0), ACCENT),
        ("TOPPADDING", (0, 0), (-1, 0), 7),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 7),
        ("TOPPADDING", (0, 1), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 1), (-1, -1), 0.5, RULE),
    ]
    for idx in range(2, len(rows), 2):
        table_style.append(("BACKGROUND", (0, idx), (-1, idx), BAND))
    table.setStyle(TableStyle(table_style))
    flow.append(table)
    flow.append(Spacer(1, 10))

    # Totals — tax row only when there is tax
    totals = [["Subtotal", fmt(invoice.get("subtotal", 0))]]
    if D(invoice.get("tax_rate", 0)) > 0:
        rate = D(invoice.get("tax_rate", 0)).normalize()
        totals.append([f"Tax ({rate:g}%)", fmt(invoice.get("tax_amount", 0))])
    totals.append(["TOTAL DUE", fmt(invoice.get("total", 0))])

    totals_rows = []
    for i, (label, value) in enumerate(totals):
        last = i == len(totals) - 1
        lstyle = ParagraphStyle(
            f"tl{i}", parent=st_td, alignment=TA_RIGHT,
            fontName="Helvetica-Bold" if last else "Helvetica",
            fontSize=11 if last else 9.5,
            textColor=colors.white if last else INK)
        vstyle = ParagraphStyle(
            f"tv{i}", parent=lstyle, fontName="Helvetica-Bold",
            fontSize=11 if last else 9.5)
        totals_rows.append([Paragraph(label, lstyle), Paragraph(value, vstyle)])

    totals_table = Table(totals_rows, colWidths=[W * 0.33, W * 0.22], hAlign="RIGHT")
    totals_table.setStyle(TableStyle([
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("LINEABOVE", (0, 0), (-1, 0), 0.5, RULE),
        ("BACKGROUND", (0, len(totals_rows) - 1), (-1, len(totals_rows) - 1), ACCENT),
    ]))
    flow.append(totals_table)
    flow.append(Spacer(1, 18))

    # Payment block — kept on one page with its heading
    payment = invoice.get("payment_details") or freelancer.get("payment_details") or ""
    if payment:
        block = [Paragraph("PAYMENT DETAILS", st_label),
                 Paragraph(_clean_inline(payment).replace("\n", "<br/>"), st_note)]
        flow.append(KeepTogether(block))
        flow.append(Spacer(1, 10))

    if invoice.get("notes"):
        flow.append(KeepTogether([
            Paragraph("NOTES", st_label),
            Paragraph(_clean_inline(invoice["notes"]), st_note)]))
        flow.append(Spacer(1, 10))

    flow.append(Spacer(1, 6))
    flow.append(Table([[""]], colWidths=[W],
                      style=TableStyle([("LINEBELOW", (0, 0), (-1, -1), 0.5, RULE)])))
    flow.append(Spacer(1, 6))
    flow.append(Paragraph(
        f"Payment due by {pretty_date(invoice.get('due_date'))}. "
        f"Please reference invoice #{invoice.get('number')} with your payment.", st_foot))

    doc.build(flow)
    return buf.getvalue()


def invoice_preview_text(invoice: dict) -> str:
    """Monospaced chat preview. Tax line appears only when tax applies."""
    lines = [
        f"Invoice #{invoice.get('number')} — {invoice.get('client_name', '')}",
        f"Project: {invoice.get('project', '—')}",
        f"Issued: {pretty_date(invoice.get('issue_date'))}   Due: {pretty_date(invoice.get('due_date'))}",
        "",
    ]
    for item in invoice.get("line_items", []):
        desc = item.get("description", "")
        if len(desc) > 42:
            desc = desc[:39] + "..."
        lines.append(
            f"  {desc:<42} {fmt_hours(item.get('hours', 0)):>6} hrs "
            f"x {fmt(item.get('rate', 0)):>10} {fmt(item.get('subtotal', 0)):>12}")

    lines.append("")
    lines.append(f"  {'Subtotal':<62}{fmt(invoice.get('subtotal', 0)):>12}")
    if D(invoice.get("tax_rate", 0)) > 0:
        rate = D(invoice.get("tax_rate", 0)).normalize()
        lines.append(f"  {f'Tax ({rate:g}%)':<62}{fmt(invoice.get('tax_amount', 0)):>12}")
    lines.append(f"  {'TOTAL DUE':<62}{fmt(invoice.get('total', 0)):>12}")
    return "\n".join(lines)
